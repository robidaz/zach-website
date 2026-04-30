from __future__ import annotations

import argparse
import re
from zipfile import ZipFile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:
    RapidOCR = None


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

HEADING_TOKEN = "@@HEADING@@"
NORMALIZED_IMAGE_MAX_WIDTH = 900
MIN_OCR_TEXT_CHARS = 4
MIN_OCR_CONFIDENCE = 0.35
MIN_DETECTED_TEXT_HEIGHT_PX = 6
STANDARD_MARKDOWN_FONT_PX = 12

SECTION_REFERENCE_RE = re.compile(r"\bsection\s+([0-9]+(?:\.[0-9]+)*)\b", re.IGNORECASE)
SECTION_SYMBOL_RE = re.compile(r"§\s*([0-9]+(?:\.[0-9]+)*)")
SEE_REFERENCE_RE = re.compile(r"\b([Ss]ee)\s+([0-9]+(?:\.[0-9]+)*)\b")
APPENDIX_REFERENCE_RE = re.compile(r"\bAppendix\s+([A-Z])\b")
MARKDOWN_LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\([^)]+\)")

_OCR_ENGINE = None


class ImageStore:
    def __init__(self, media_dir: Path) -> None:
        self.media_dir = media_dir
        self.media_dir.mkdir(parents=True, exist_ok=True)
        self._saved: dict[str, str] = {}
        self._text_image_widths: dict[str, int | None] = {}

    def save(self, rel_id: str, part) -> str:
        saved = self._saved.get(rel_id)
        if saved:
            return saved

        image_part = part.related_parts[rel_id]
        file_name = Path(image_part.partname).name
        if Path(file_name).suffix.lower() != ".emf":
            target = self.media_dir / file_name
            if not target.exists():
                target.write_bytes(image_part.blob)
        self._saved[rel_id] = file_name
        return file_name

    def text_render_width(self, file_name: str) -> int | None:
        if file_name in self._text_image_widths:
            return self._text_image_widths[file_name]

        image_path = self.media_dir / file_name
        render_width = text_render_width(image_path)
        self._text_image_widths[file_name] = render_width
        return render_width


def get_ocr_engine():
    global _OCR_ENGINE

    if RapidOCR is None:
        return None

    if _OCR_ENGINE is None:
        try:
            _OCR_ENGINE = RapidOCR()
        except Exception:
            _OCR_ENGINE = False

    return _OCR_ENGINE or None


def detected_text_metrics(image_path: Path) -> tuple[int, list[float]]:
    if not image_path.exists() or Image is None:
        return 0, []

    engine = get_ocr_engine()
    if engine is None:
        return 0, []

    try:
        with Image.open(image_path) as image:
            image_width, _ = image.size
        result, _ = engine(str(image_path))
    except Exception:
        return 0, []

    if not result:
        return image_width, []

    confident_tokens: list[str] = []
    heights: list[float] = []
    xs: list[float] = []
    for line in result:
        if len(line) < 3:
            continue

        box, token, confidence = line[:3]
        token = str(token).strip()
        if not token or not any(char.isalnum() for char in token):
            continue

        try:
            confidence_value = float(confidence)
        except (TypeError, ValueError):
            confidence_value = -1

        if confidence_value >= MIN_OCR_CONFIDENCE:
            confident_tokens.append(token)

            try:
                x_positions = [float(point[0]) for point in box]
                ys = [float(point[1]) for point in box]
                height = max(ys) - min(ys)
            except (TypeError, ValueError, IndexError):
                continue

            if height >= MIN_DETECTED_TEXT_HEIGHT_PX:
                heights.append(height)
                xs.extend([min(x_positions), max(x_positions)])

    if len("".join(confident_tokens)) < MIN_OCR_TEXT_CHARS:
        return image_width, []

    if xs:
        image_width = round(max(xs) - min(xs))

    return image_width, heights


def text_render_width(image_path: Path) -> int | None:
    image_width, heights = detected_text_metrics(image_path)
    if not image_width or not heights:
        return None

    smallest_text_height = min(heights)
    scale_ratio = STANDARD_MARKDOWN_FONT_PX / smallest_text_height
    return max(1, round(image_width * scale_ratio))


def image_markup(file_name: str, image_store: ImageStore) -> str:
    if Path(file_name).suffix.lower() == ".emf":
        return (
            '<div style="height: 2.5rem; border: 1px solid #999; border-radius: 4px;"></div>'
            '<div style="margin-top: 0.35rem; border-top: 1px solid #000;"></div>'
        )

    render_width = image_store.text_render_width(file_name)
    if render_width is not None:
        return (
            f'<img src="media/{file_name}" alt="{file_name}" '
            f'style="display: block; width: {render_width}px; max-width: 100%; height: auto;" />'
        )

    return f"![{file_name}](media/{file_name})"


def iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_images(element, part, image_store: ImageStore) -> list[str]:
    image_refs: list[str] = []
    rel_ids: list[str] = []

    for blip in element.findall(".//a:blip", NS):
        rel_id = blip.get(qn("r:embed"))
        if rel_id:
            rel_ids.append(rel_id)

    for imagedata in element.findall(".//v:imagedata", NS):
        rel_id = imagedata.get(qn("r:id"))
        if rel_id:
            rel_ids.append(rel_id)

    for rel_id in rel_ids:
        file_name = image_store.save(rel_id, part)
        image_refs.append(image_markup(file_name, image_store))

    return image_refs


def extract_textboxes(element) -> list[str]:
    blocks: list[str] = []
    for textbox in element.findall(".//w:txbxContent", NS):
        paragraphs: list[str] = []
        for paragraph in textbox.findall(".//w:p", NS):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()
            if text:
                paragraphs.append(text)
        if paragraphs:
            block = "\n".join(paragraphs)
            if not blocks or blocks[-1] != block:
                blocks.append(block)
    return blocks


def extract_footer_text(input_path: Path) -> list[str]:
    footer_lines: list[str] = []
    with ZipFile(input_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/footer") or not name.endswith(".xml"):
                continue
            root = archive.read(name)
            for text in re.findall(rb">([^<]+)<", root):
                decoded = text.decode("utf-8", errors="ignore").strip()
                if decoded and decoded not in footer_lines:
                    footer_lines.append(decoded)
    return footer_lines


def extract_run_text(run_element) -> str:
    fragments: list[str] = []
    for child in run_element.iterchildren():
        if child.tag == qn("w:t"):
            fragments.append(child.text or "")
        elif child.tag == qn("w:tab"):
            fragments.append("\t")
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            fragments.append("\n")
        elif child.tag == qn("w:noBreakHyphen"):
            fragments.append("-")
    return "".join(fragments)


def extract_hyperlink_text(hyperlink_element, part, image_store: ImageStore) -> tuple[str, list[str]]:
    text_parts: list[str] = []
    image_refs: list[str] = []

    for child in hyperlink_element.iterchildren():
        if child.tag == qn("w:r"):
            text_parts.append(extract_run_text(child))
            image_refs.extend(extract_images(child, part, image_store))

    text = "".join(text_parts)
    rel_id = hyperlink_element.get(qn("r:id"))
    if rel_id and text:
        target = part.rels[rel_id].target_ref
        text = f"[{text}]({target})"
    return text, image_refs


def extract_paragraph_content(paragraph: Paragraph, image_store: ImageStore) -> tuple[str, list[str]]:
    text_parts: list[str] = []
    image_refs: list[str] = []

    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            text_parts.append(extract_run_text(child))
            text_parts.extend(extract_textboxes(child))
            image_refs.extend(extract_images(child, paragraph.part, image_store))
        elif child.tag == qn("w:hyperlink"):
            hyperlink_text, hyperlink_images = extract_hyperlink_text(child, paragraph.part, image_store)
            text_parts.append(hyperlink_text)
            text_parts.extend(extract_textboxes(child))
            image_refs.extend(hyperlink_images)

    return "".join(text_parts).strip(), image_refs


def heading_spec(style_name: str) -> tuple[str, int] | None:
    if style_name == "Title":
        return "title", 1
    if style_name == "Heading - Non-Numbered":
        return "plain", 1
    if style_name == "Heading 1 - Alphabetical":
        return "appendix", 1
    match = re.match(r"Heading\s+(\d+)", style_name)
    if match:
        return "numeric", max(1, min(6, int(match.group(1))))
    return None


def slugify_heading(text: str) -> str:
    # Matches markdownlint/GitHub slug algorithm: keep only [a-z0-9 _-], then
    # replace each individual space with '-' (no run-collapsing, no hyphen-collapsing).
    slug = text.casefold()
    slug = re.sub(r"[^a-z0-9 _-]", "", slug)
    slug = slug.replace("_", "")
    slug = slug.strip().replace(" ", "-")
    return slug


def build_toc_block(headings: list[tuple[int, str, str]]) -> str:
    lines = ["## Table of Contents", "", "<!-- TOC -->"]
    for level, text, anchor in headings:
        indent = "  " * (level - 1)
        lines.append(f"{indent}- [{text}](#{anchor})")
    lines.append("<!-- /TOC -->")
    return "\n".join(lines)


def linkify_block_references(
    block: str,
    section_anchor_map: dict[str, str],
    appendix_anchor_map: dict[str, str],
) -> str:
    if block.startswith("#"):
        return block

    placeholders: list[str] = []

    def protect_link(match: re.Match[str]) -> str:
        placeholders.append(match.group(0))
        return f"@@LINK{len(placeholders) - 1}@@"

    def restore_links(value: str) -> str:
        for index, link in enumerate(placeholders):
            value = value.replace(f"@@LINK{index}@@", link)
        return value

    block = MARKDOWN_LINK_RE.sub(protect_link, block)

    def replace_section(match: re.Match[str]) -> str:
        section_number = match.group(1)
        anchor = section_anchor_map.get(section_number)
        if anchor is None:
            return match.group(0)
        return f"[{match.group(0)}](#{anchor})"

    def replace_symbol(match: re.Match[str]) -> str:
        section_number = match.group(1)
        anchor = section_anchor_map.get(section_number)
        if anchor is None:
            return match.group(0)
        return f"[§{section_number}](#{anchor})"

    def replace_see(match: re.Match[str]) -> str:
        lead = match.group(1)
        section_number = match.group(2)
        anchor = section_anchor_map.get(section_number)
        if anchor is None:
            return match.group(0)
        return f"{lead} [{section_number}](#{anchor})"

    def replace_appendix(match: re.Match[str]) -> str:
        appendix_letter = match.group(1)
        anchor = appendix_anchor_map.get(appendix_letter)
        if anchor is None:
            return match.group(0)
        return f"[Appendix {appendix_letter}](#{anchor})"

    block = SECTION_REFERENCE_RE.sub(replace_section, block)
    block = SECTION_SYMBOL_RE.sub(replace_symbol, block)
    block = SEE_REFERENCE_RE.sub(replace_see, block)
    block = APPENDIX_REFERENCE_RE.sub(replace_appendix, block)
    return restore_links(block)


def finalize_markdown(blocks: list[str]) -> str:
    counters = [0] * 6
    appendix_index = 0
    title_seen = False
    first_content_heading_index: int | None = None
    toc_headings: list[tuple[int, str, str]] = []
    section_anchor_map: dict[str, str] = {}
    appendix_anchor_map: dict[str, str] = {}
    processed_blocks: list[str] = []

    for block in blocks:
        lines = block.split("\n")
        first_line = lines[0] if lines else ""
        if not first_line.startswith(f"{HEADING_TOKEN}\t"):
            processed_blocks.append(block)
            continue

        _, heading_kind, level_text, raw_heading_text = first_line.split("\t", 3)
        level = int(level_text)
        trailing_lines = [line for line in lines[1:] if line]
        include_in_toc = False

        if heading_kind == "title":
            heading_text = raw_heading_text
            title_seen = True
        elif heading_kind == "plain":
            heading_text = raw_heading_text
            include_in_toc = True
        elif heading_kind == "appendix":
            appendix_index += 1
            appendix_letter = chr(ord("A") + appendix_index - 1)
            appendix_title = re.sub(r"^Appendix:\s*", "", raw_heading_text).strip()
            heading_text = f"Appendix {appendix_letter}: {appendix_title}"
            appendix_anchor_map[appendix_letter] = slugify_heading(heading_text)
            include_in_toc = True
        else:
            counters[level - 1] += 1
            for index in range(level, len(counters)):
                counters[index] = 0
            section_number = ".".join(str(value) for value in counters[:level] if value)
            heading_text = f"{section_number}. {raw_heading_text}"
            section_anchor_map[section_number] = slugify_heading(heading_text)
            include_in_toc = True

        if heading_kind not in ("title", "plain") and first_content_heading_index is None:
            first_content_heading_index = len(processed_blocks)

        heading_block = "\n".join([f"{'#' * level} {heading_text}", *trailing_lines]).strip()
        if include_in_toc:
            toc_headings.append((level, heading_text, slugify_heading(heading_text)))
        processed_blocks.append(heading_block)

    processed_blocks = [
        linkify_block_references(block, section_anchor_map, appendix_anchor_map)
        for block in processed_blocks
    ]

    if toc_headings:
        toc_block = build_toc_block(toc_headings)
        insert_at = first_content_heading_index if first_content_heading_index is not None else 0
        processed_blocks.insert(insert_at, toc_block)

    return "\n\n".join(processed_blocks).strip() + "\n"


def rebuild_existing_markdown(input_path: Path, markdown_path: Path) -> None:
    document = Document(str(input_path))
    heading_meta: list[tuple[str, int, str]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading = heading_spec(style_name)
        if heading is None:
            continue
        heading_kind, level = heading
        heading_meta.append((heading_kind, level, text))

    markdown_text = MARKDOWN_LINK_RE.sub(r"\1", markdown_path.read_text(encoding="utf-8"))
    blocks = re.split(r"\n\s*\n", markdown_text.strip())
    converted_blocks: list[str] = []
    heading_index = 0
    title_preserved = False
    for block in blocks:
        if block.startswith("## Table of Contents") or block.startswith("<!-- TOC -->"):
            continue

        lines = block.splitlines()
        if lines and re.match(r"^#{1,6} ", lines[0]) and heading_index < len(heading_meta):
            current_heading_text = re.sub(r"^#{1,6}\s+", "", lines[0]).strip()
            if not title_preserved and current_heading_text != heading_meta[heading_index][2]:
                converted_blocks.append(
                    "\n".join([f"{HEADING_TOKEN}\ttitle\t1\t{current_heading_text}", *lines[1:]]).strip()
                )
                title_preserved = True
                continue

            heading_kind, level, heading_text = heading_meta[heading_index]
            converted_blocks.append(
                "\n".join([f"{HEADING_TOKEN}\t{heading_kind}\t{level}\t{heading_text}", *lines[1:]]).strip()
            )
            heading_index += 1
        else:
            converted_blocks.append(block)

    markdown_path.write_text(finalize_markdown(converted_blocks), encoding="utf-8")


def list_marker(style_name: str) -> tuple[str, int] | None:
    if style_name == "Bulleted List" or style_name == "List Paragraph":
        return "-", 0

    bullet_match = re.match(r"List Bullet(?:\s+(\d+))?$", style_name)
    if bullet_match:
        level = int(bullet_match.group(1) or "1") - 1
        return "-", max(0, level)

    number_match = re.match(r"List Number(?:\s+(\d+))?$", style_name)
    if number_match:
        level = int(number_match.group(1) or "1") - 1
        return "1.", max(0, level)

    return None


def render_paragraph(paragraph: Paragraph, image_store: ImageStore) -> str:
    text, image_refs = extract_paragraph_content(paragraph, image_store)
    style_name = paragraph.style.name if paragraph.style is not None else ""

    lines: list[str] = []
    if text:
        heading = heading_spec(style_name)
        marker = list_marker(style_name)
        if heading is not None:
            heading_kind, level = heading
            lines.append(f"{HEADING_TOKEN}\t{heading_kind}\t{level}\t{text}")
        elif marker is not None:
            bullet, indent_level = marker
            indent = "  " * indent_level
            lines.append(f"{indent}{bullet} {text}")
        else:
            lines.append(text)

    lines.extend(image_refs)
    return "\n".join(lines).strip()


def escape_table_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", "<br>")


def render_cell(cell: _Cell, image_store: ImageStore) -> str:
    parts: list[str] = []
    for block in iter_block_items(cell):
        if isinstance(block, Paragraph):
            text, image_refs = extract_paragraph_content(block, image_store)
            if text:
                parts.append(text)
            parts.extend(image_refs)
        elif isinstance(block, Table):
            nested = render_table(block, image_store)
            if nested:
                parts.append(nested)
    while parts and not parts[-1]:
        parts.pop()
    return escape_table_cell("<br><br>".join(parts))


def render_html_table(rows: list[list[str]]) -> str:
    lines = ["<table>"]
    for row in rows:
        lines.append("  <tr>")
        for cell in row:
            lines.append(f"    <td>{cell}</td>")
        lines.append("  </tr>")
    lines.append("</table>")
    return "\n".join(lines)


def render_table(table: Table, image_store: ImageStore) -> str:
    rows = [[render_cell(cell, image_store) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]

    if len(normalized_rows) == 1:
        return render_html_table(normalized_rows)

    header = normalized_rows[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * column_count) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert(input_path: Path, output_path: Path) -> None:
    document = Document(str(input_path))
    image_store = ImageStore(output_path.parent / "media")

    blocks: list[str] = []
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            rendered = render_paragraph(block, image_store)
        else:
            rendered = render_table(block, image_store)

        if rendered:
            blocks.append(rendered)

    footer_lines = extract_footer_text(input_path)
    if footer_lines:
        blocks.append("\n".join(footer_lines))

    output_path.write_text(finalize_markdown(blocks), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path, nargs="?")
    parser.add_argument("--existing-markdown", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = args.input.resolve()
    if args.existing_markdown is not None:
        markdown_path = args.existing_markdown.resolve()
        rebuild_existing_markdown(input_path, markdown_path)
        print(f"Updated: {markdown_path}")
        return

    output_path = (args.output or input_path.with_suffix(".md")).resolve()
    convert(input_path, output_path)
    print(f"Converted: {output_path}")


if __name__ == "__main__":
    main()
